"""
Generator dokumentów .docx dla systemu obsługi praktyk zawodowych IIS ANS Elbląg.
Odwzorowuje formatowanie oryginałów Word 1:1.
"""
from io import BytesIO
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

FONT_NAME = "Cambria"
FONT_BODY = Pt(11)
FONT_SM   = Pt(10)
FONT_XS   = Pt(9)

SURVEY_QUESTIONS = [
    "Poznałam/poznałem zasady funkcjonowania instytucji, w której odbywałam/odbywałem praktyki zawodowe.",
    "Poznałam/poznałem strukturę oraz regulamin organizacyjny instytucji, w której odbywałam/odbywałem praktyki zawodowe.",
    "Praktyki zawodowe umożliwiły mi pełną realizację ramowego programu praktyk zawodowych przewidzianego w ramach mojego kierunku studiów.",
    "Podczas praktyk zawodowych zwracano uwagę na przestrzeganie zasad etyki i tajemnicy zawodowej.",
    "Podczas praktyk miałam/miałem możliwość praktycznego zastosowania wiedzy teoretycznej zdobytej na zajęciach.",
    "Praktyki zawodowe przyczyniły się do pogłębienia mojej wiedzy i umiejętności zdobytych w trakcie studiów.",
    "Mogłem liczyć na wsparcie merytoryczne Opiekuna zakładowego praktyk.",
    "Mogłem liczyć na wsparcie merytoryczne Opiekuna uczelnianego praktyk.",
    "Opiekun zakładowy odpowiedzialny za praktyki zawodowe w miejscu ich odbywania potrafił prawidłowo zorganizować ich przebieg.",
    "Podczas praktyk zawodowych miałam/miałem możliwość pozyskiwania materiałów niezbędnych do przygotowania mojej pracy dyplomowej.",
    "Praktyki zawodowe rozwinęły moje umiejętności skutecznego komunikowania się w sytuacjach zawodowych i pracy w zespole.",
    "Praktyki zawodowe nauczyły mnie samodzielności i odpowiedzialności podczas wykonywania pracy.",
    "Liczba godzin realizowana w ramach praktyk zawodowych jest wystarczająca.",
    "Czy po zakończeniu praktyki zawodowej chciałaby/chciałby Pani/Pan współpracować z instytucją, w której Pani/Pan zrealizowała/zrealizował praktykę?",
]

LEARNING_EFFECTS = [
    "Ma wiedzę na temat sposobu realizacji zadań inżynierskich dotyczących informatyki z zachowaniem standardów i norm technicznych.",
    "Zna technologie, narzędzia, metody, techniki oraz sprzęt stosowane w informatyce.",
    "Zna ekonomiczne, prawne skutki własnych działań podejmowanych w ramach praktyki oraz ograniczenia wynikające z prawa autorskiego.",
    "Zna zasady bezpieczeństwa pracy i ergonomii w zawodzie informatyka.",
    "Pozyskuje informacje odnośnie technologii, metod, technik, sprzętu wymaganego do realizacji powierzonego zadania.",
    "W oparciu o kontakty ze środowiskiem inżynierskim zakładu, potrafi podnieść swoje umiejętności zawodowe.",
    "Opracowuje dokumentację dotyczącą realizacji podejmowanych zadań w ramach praktyki.",
    "Potrafi zidentyfikować problem informatyczny występujący w zakładzie pracy / instytucji.",
    "Potrafi rozwiązać rzeczywiste zadanie inżynierskie z zakresu działalności informatycznej.",
    "Pracuje w zespole zajmującym się zawodowo branżą IT.",
    "Przestrzega zasad etyki zawodowej i zgodnie z tymi zasadami korzysta z wiedzy i pomocy doświadczonych kolegów.",
    "Kontaktując się z osobami spoza branży potrafi zarówno pozyskać od nich niezbędne informacje jak i przekazać własne.",
    "Dostrzega w praktyce tempo deaktualizacji wiedzy informatycznej oraz skutki działalności informatyków.",
]


# ─── helpers ──────────────────────────────────────────────────────────────────

def _new_doc(top=20, bot=20, left=25, right=25):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin    = Mm(top)
    sec.bottom_margin = Mm(bot)
    sec.left_margin   = Mm(left)
    sec.right_margin  = Mm(right)
    # Remove default empty paragraph
    for p in doc.paragraphs:
        _remove_para(p)
    return doc


def _remove_para(p):
    el = p._element
    el.getparent().remove(el)


def _run(para, text, bold=False, italic=False, size=None, font=FONT_NAME, underline=False):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = size or FONT_BODY
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    return run


def _para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
          size=None, italic=False, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if text:
        _run(p, text, bold=bold, size=size, italic=italic)
    return p


def _para_center(doc, text, bold=False, size=None, space_before=0, space_after=4):
    return _para(doc, text, WD_ALIGN_PARAGRAPH.CENTER, bold=bold, size=size,
                 space_before=space_before, space_after=space_after)


def _set_cell_bg(cell, color_hex="F2F2F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, top=True, bot=True, left=True, right=True, color="000000"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, active in [("top", top), ("bottom", bot), ("left", left), ("right", right)]:
        border = OxmlElement(f"w:{side}")
        if active:
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), color)
        else:
            border.set(qn("w:val"), "none")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _cell_para(cell, text="", bold=False, size=None, italic=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if text:
        _run(p, text, bold=bold, size=size, italic=italic)
    return p


def _cell_add_para(cell, text="", bold=False, size=None, italic=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2):
    p = cell.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if text:
        _run(p, text, bold=bold, size=size, italic=italic)
    return p


def _set_col_widths(table, widths_mm):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_mm):
                cell.width = Mm(widths_mm[j])


def _letterhead(doc, badge_nr):
    """Nagłówek ANS z numerem załącznika – ta sama tabela co w oryginałach."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    # Hide all borders
    for cell in tbl.rows[0].cells:
        _set_cell_borders(cell, False, False, False, False)

    left_cell  = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]
    left_cell.width  = Mm(130)
    right_cell.width = Mm(30)

    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(lp, "Akademia Nauk Stosowanych w Elblągu", bold=True, size=Pt(11))
    _cell_add_para(left_cell, "Instytut Informatyki Stosowanej im. Krzysztofa Brzeskiego",
                   bold=False, size=Pt(10))

    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(rp, f"Załącznik nr {badge_nr}", bold=False, size=Pt(9))

    _para(doc, "", space_before=0, space_after=4)


def _doc_title(doc, title, subtitle=None, center=True):
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p = _para(doc, title, align=align, bold=True, size=Pt(13), space_before=4, space_after=2)
    if subtitle:
        _para(doc, subtitle, align=align, bold=False, size=Pt(11), space_before=0, space_after=6)
    else:
        _para(doc, "", space_before=0, space_after=4)
    return p


def _section_label(doc, text):
    p = _para(doc, text, bold=True, size=Pt(11), space_before=6, space_after=2)
    # underline the section label
    for run in p.runs:
        run.underline = True
    return p


def v(data, key):
    """Safely get value from data dict, return empty string if missing/None."""
    if not data:
        return ""
    val = data.get(key)
    if val is None:
        return ""
    return str(val)


def _field_row_2col(doc, label1, val1, label2, val2, w1=80, w2=80):
    """Two labeled fields side by side in a borderless table."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in tbl.rows:
        for cell in row.cells:
            _set_cell_borders(cell, False, False, False, False)
    _set_col_widths(tbl, [w1, w2])

    _cell_para(tbl.rows[0].cells[0], label1, bold=False, size=Pt(9),
               align=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)
    _cell_para(tbl.rows[0].cells[1], label2, bold=False, size=Pt(9),
               align=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)

    c00 = tbl.rows[1].cells[0]
    c01 = tbl.rows[1].cells[1]
    _cell_para(c00, val1 or "—", bold=False, size=FONT_BODY, space_after=2)
    _cell_para(c01, val2 or "—", bold=False, size=FONT_BODY, space_after=2)
    # Bottom border = underline effect
    _set_cell_borders(c00, top=False, bot=True, left=False, right=False)
    _set_cell_borders(c01, top=False, bot=True, left=False, right=False)
    _para(doc, "", space_before=0, space_after=3)


def _field_single(doc, label, val, width_mm=160):
    p_label = _para(doc, label, size=Pt(9), space_before=4, space_after=1)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Mm(width_mm)
    _set_cell_borders(cell, top=False, bot=True, left=False, right=False)
    _cell_para(cell, val or "—", size=FONT_BODY, space_after=2)
    _para(doc, "", space_before=0, space_after=3)


def _sig_row(doc, label1, val1, label2=None, val2=None):
    cols = 2
    tbl = doc.add_table(rows=2, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in tbl.rows:
        for cell in row.cells:
            _set_cell_borders(cell, False, False, False, False)

    w = 80
    _set_col_widths(tbl, [w, w])

    _cell_para(tbl.rows[0].cells[0], label1, bold=False, size=Pt(9), space_after=1)
    if label2:
        _cell_para(tbl.rows[0].cells[1], label2, bold=False, size=Pt(9), space_after=1)

    c0 = tbl.rows[1].cells[0]
    c1 = tbl.rows[1].cells[1]
    _cell_para(c0, val1 or "—", size=FONT_BODY, space_after=2)
    _set_cell_borders(c0, top=False, bot=True, left=False, right=False)
    if val2 is not None:
        _cell_para(c1, val2 or "—", size=FONT_BODY, space_after=2)
        _set_cell_borders(c1, top=False, bot=True, left=False, right=False)

    _para(doc, "", space_before=0, space_after=3)


def _to_bytes(doc):
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── Zał. 1 – Porozumienie z zakładem pracy ───────────────────────────────────

def gen_zal1(data):
    doc = _new_doc()
    _letterhead(doc, "1")
    _doc_title(doc, "Porozumienie o organizację praktyki zawodowej\nstudenta Akademii Nauk Stosowanych w Elblągu")

    # Dane dokumentu
    _section_label(doc, "Dane dokumentu")
    p = _para(doc, "", space_before=2, space_after=2)
    _run(p, "Porozumienie nr ")
    _run(p, v(data, "nr_porozumienia") or "……………", bold=True)
    _run(p, "  zawarte w  ")
    _run(p, v(data, "miejscowosc") or "………………………", bold=True)
    _run(p, "  dnia  ")
    _run(p, v(data, "data") or "………………………", bold=True)

    # A. Dane studenta
    _section_label(doc, "A. Dane studenta")
    _field_row_2col(doc, "Imię i nazwisko", v(data, "imie_nazwisko"),
                    "Nr albumu", v(data, "nr_albumu"), w1=105, w2=55)
    _field_row_2col(doc, "Kierunek studiów", "Informatyka",
                    "Specjalność", v(data, "specjalnosc"), w1=55, w2=105)
    _field_single(doc, "Forma studiów", v(data, "rodzaj_studiow") or "stacjonarne")

    # B. Dane zakładu pracy
    _section_label(doc, "B. Dane zakładu pracy")
    _field_single(doc, "Nazwa instytucji / zakładu pracy", v(data, "nazwa_zakladu"))
    _field_row_2col(doc, "Adres", v(data, "adres_zakladu"),
                    "NIP", v(data, "nip_zakladu"), w1=105, w2=55)
    _field_row_2col(doc, "Reprezentant – imię i nazwisko", v(data, "reprezentant_nazwisko"),
                    "Stanowisko / funkcja", v(data, "reprezentant_stanowisko"))

    # C. Termin praktyki
    _section_label(doc, "C. Termin praktyki")
    _field_row_2col(doc, "Data rozpoczęcia", v(data, "data_start"),
                    "Data zakończenia", v(data, "data_end"))
    _field_row_2col(doc, "Liczba godzin", v(data, "liczba_godzin") or "960",
                    "Uczelniany opiekun praktyki", v(data, "uczelniany_opiekun"))

    # D. Podpisy
    _section_label(doc, "D. Podpisy")
    _sig_row(doc,
             "Podpis zakładowego opiekuna", v(data, "podpis_zakladowy"),
             "Podpis uczelnianego opiekuna / Dyrektor IIS", v(data, "podpis_uczelniany"))

    return _to_bytes(doc)


# ─── Zał. 2 – Program praktyki zawodowej ──────────────────────────────────────

def gen_zal2(data):
    doc = _new_doc()
    _letterhead(doc, "2")
    _doc_title(doc, "Program praktyki zawodowej")

    _para_center(doc, "A.  Etap pierwszy – rozpoczęcie praktyki", bold=True, size=Pt(11), space_before=4)
    items_a = [
        "Czynności organizacyjne, szkolenie BHP i ppoż.",
        "Zapoznanie się z zakresem działania zakładu pracy (rodzaj usług, produkcja, struktura organizacyjna itp.) ze szczególnym uwzględnieniem infrastruktury IT.",
        "Zapoznanie z projektami realizowanymi przez firmę, środkami produkcji, aplikacjami użytkowymi, stosowanymi technologiami i metodami pracy.",
    ]
    for item in items_a:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        _run(p, item, size=FONT_BODY)

    _para_center(doc, "B.  Etap drugi", bold=True, size=Pt(11), space_before=6)
    _para(doc, "Praca na ostatecznym stanowisku pracy, wykonywanie prac i projektów informatycznych tak aby osiągnąć poniższe, wymagane efekty uczenia się:", size=FONT_BODY, space_after=2)

    for i, eff in enumerate(LEARNING_EFFECTS, 1):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        _run(p, eff, size=FONT_BODY)

    _para(doc, "Opcjonalnie formułowanie przez zakład pracy tematu pracy dyplomowej i precyzowanie zakresu pracy dyplomowej; rozpoczęcie pracy nad projektem dyplomowym.", size=FONT_BODY, space_before=4, space_after=2)
    _para(doc, "Planowany czas realizacji praktyki: 6 miesięcy tj. 120 dni (960 godz.).", bold=True, size=FONT_BODY, space_after=4)

    _para_center(doc, "C.  Etap trzeci - zakończenie praktyki", bold=True, size=Pt(11), space_before=4)
    items_c = [
        'W trakcie praktyki student prowadzi "Dzienniczek praktyki". Wpisy w dzienniczku powinny być potwierdzone przez zakładowego opiekuna praktyki.',
        'Na "Karcie praktyki zawodowej" kierownik zakładu pracy potwierdza odbycie praktyki a opiekun zakładowy wystawia ocenę parametryczną i opisową.',
        'Opiekun ze strony zakładu pracy na odpowiednim druku "Potwierdzenie efektów uczenia się" potwierdza osiągnięcie przez praktykanta zakładanych efektów uczenia się.',
    ]
    for item in items_c:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        _run(p, item, size=FONT_BODY)

    _para(doc, "", space_before=8, space_after=2)
    _sig_row(doc,
             "Zakład pracy", v(data, "podpis_zakladowy"),
             "Dyrektor Instytutu", v(data, "podpis_uczelniany"))

    return _to_bytes(doc)


# ─── Zał. 2a – Program i harmonogram ──────────────────────────────────────────

def gen_zal2a(data, effects=None):
    doc = _new_doc(top=20, bot=25, left=15, right=15)
    _letterhead(doc, "2a")
    _doc_title(doc, "Program i harmonogram praktyki zawodowej")

    _field_row_2col(doc, "Student/ka", v(data, "imie_nazwisko"),
                    "Nr albumu", v(data, "nr_albumu"), w1=110, w2=60)
    _field_row_2col(doc, "Kierunek studiów", "Informatyka",
                    "Specjalność", v(data, "specjalnosc"), w1=60, w2=110)
    _field_single(doc, "Miejsce praktyki (instytucja)", v(data, "miejsce_praktyki"))
    _field_row_2col(doc, "Termin realizacji praktyki: od", v(data, "data_start"),
                    "do", v(data, "data_end"), w1=80, w2=90)

    # Tabela efektów + dział
    _section_label(doc, "Program praktyki zawodowej")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    _cell_para(hdr[0], "Nr", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr[1], "Efekty kształcenia", bold=True, size=Pt(9))
    _cell_para(hdr[2], "Dział (komórka) / przykładowe prace wykonywane przez praktykanta",
               bold=True, size=Pt(9))
    _set_cell_bg(hdr[0])
    _set_cell_bg(hdr[1])
    _set_cell_bg(hdr[2])

    efekty_plan = {e["nr"]: e.get("dzial_prace", "") for e in (data.get("efekty_plan") or [])}
    effect_list = effects or []
    for eff in effect_list:
        row = tbl.add_row()
        _cell_para(row.cells[0], f"{eff.nr:02d}", size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], eff.opis, size=Pt(9))
        _cell_para(row.cells[2], efekty_plan.get(eff.nr, ""), size=FONT_BODY)
    _set_col_widths(tbl, [12, 85, 83])

    # Harmonogram
    _section_label(doc, "Harmonogram praktyki zawodowej")
    htbl = doc.add_table(rows=1, cols=4)
    htbl.style = "Table Grid"
    htbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hh = htbl.rows[0].cells
    _cell_para(hh[0], "L.p.", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hh[1], "Dział / komórka (miejsce odbywania praktyki)", bold=True, size=Pt(9))
    _cell_para(hh[2], "Przykładowe prace", bold=True, size=Pt(9))
    _cell_para(hh[3], "Planowana liczba dni roboczych", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in hh:
        _set_cell_bg(c)

    harmono = data.get("harmonogram") or []
    for entry in harmono:
        row = htbl.add_row()
        _cell_para(row.cells[0], str(entry.get("lp", "")), size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], entry.get("dzial", ""), size=FONT_BODY)
        _cell_para(row.cells[2], "", size=FONT_BODY)
        _cell_para(row.cells[3], str(entry.get("dni", "")), size=FONT_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_col_widths(htbl, [12, 70, 70, 28])

    _para(doc, f"Uzgodniono w dniu: {v(data, 'data_uzgodnienia') or '……………………………'}", size=FONT_BODY, space_before=6, space_after=4)

    # Podpisy 3
    tbl3 = doc.add_table(rows=2, cols=3)
    tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in tbl3.rows:
        for cell in row.cells:
            _set_cell_borders(cell, False, False, False, False)
    labels = ["Podpis uczelnianego opiekuna praktyki", "Podpis zakładowego opiekuna praktyki", "Podpis studenta"]
    vals   = [v(data, "podpis_uczelniany"), v(data, "podpis_zakladowy"), v(data, "podpis_studenta")]
    for j in range(3):
        _cell_para(tbl3.rows[0].cells[j], labels[j], size=Pt(9), space_after=1)
        c = tbl3.rows[1].cells[j]
        _cell_para(c, vals[j] or "—", size=FONT_BODY, space_after=2)
        _set_cell_borders(c, top=False, bot=True, left=False, right=False)

    return _to_bytes(doc)


# ─── Zał. 3 – Karta praktyki zawodowej ────────────────────────────────────────

def gen_zal3(data):
    doc = _new_doc()
    _letterhead(doc, "3")
    _doc_title(doc, "Karta praktyki zawodowej", "Skierowanie na praktykę")

    # A – Dane skierowania
    _section_label(doc, "A. Dane skierowania")
    p = _para(doc, "", space_before=2, space_after=2)
    _run(p, "Na podstawie porozumienia nr  ")
    _run(p, v(data, "nr_porozumienia") or "……………", bold=True)
    _run(p, "  z dnia  ")
    _run(p, v(data, "data_porozumienia") or "………………………", bold=True)
    _run(p, "  r., kieruję niżej wymienionego studenta na praktykę zawodową do zakładu pracy:")
    _field_single(doc, "Nazwa instytucji / zakładu pracy", v(data, "zaklad_pracy"))

    # B – Dane studenta
    _section_label(doc, "B. Dane studenta")
    _field_row_2col(doc, "Imię i nazwisko", v(data, "imie_nazwisko"),
                    "Numer albumu", v(data, "nr_albumu"), w1=105, w2=55)
    _field_row_2col(doc, "Kierunek studiów", "Informatyka",
                    "Specjalność", v(data, "specjalnosc"), w1=55, w2=105)
    _field_single(doc, "Forma studiów", v(data, "rodzaj_studiow") or "stacjonarne")
    _para(doc, "Czas trwania praktyki: 6 miesięcy (120 dni roboczych) – inżynierskie stacjonarne",
          size=Pt(10), italic=True, space_before=2, space_after=4)
    _field_single(doc, "Uczelniany opiekun praktyki zawodowej", v(data, "uczelniany_opiekun"))
    _field_row_2col(doc, "Data rozpoczęcia praktyki", v(data, "data_start"),
                    "Data zakończenia praktyki", v(data, "data_end"))

    # C – Zakładowy opiekun
    _section_label(doc, "C. Zakładowy opiekun praktyki")
    _field_row_2col(doc, "Imię i nazwisko zakładowego opiekuna", v(data, "zakladowy_opiekun_nazwisko"),
                    "Funkcja / stanowisko", v(data, "zakladowy_opiekun_funkcja"))
    _sig_row(doc,
             "Potwierdzenie zgłoszenia studenta na praktykę", v(data, "potwierdzenie_zgloszenia"),
             "Potwierdzenie odbycia szkolenia BHP", v(data, "potwierdzenie_bhp"))

    # D – Zaświadczenie
    _section_label(doc, "D. Zaświadczenie odbycia praktyki zawodowej")
    imie = v(data, "imie_nazwisko") or "……………………………"
    p = _para(doc, f"Zaświadczam, że student {imie} odbył/-a praktykę zawodową w:", size=FONT_BODY, space_before=2, space_after=2)
    _field_single(doc, "Nazwa zakładu / instytucji", v(data, "zaswiadczenie_zaklad"))
    _field_row_2col(doc, "Okres od", v(data, "zaswiadczenie_okres_od"),
                    "Okres do", v(data, "zaswiadczenie_okres_do"))
    _field_single(doc, "Uwagi", v(data, "zaswiadczenie_uwagi"))
    _sig_row(doc, "Miejscowość, data, pieczęć i podpis kierownika zakładu", v(data, "zaswiadczenie_podpis"))

    # E – Oceny
    _section_label(doc, "E. Ocena przebiegu praktyki zawodowej")

    _para(doc, "Ocena zakładowego opiekuna praktyki", bold=True, size=FONT_BODY, space_before=4, space_after=2)
    _field_row_2col(doc, "Ocena parametryczna (skala 2–5)", v(data, "ocena_zakladowa_param") or "—",
                    "", "", w1=80, w2=80)
    _field_single(doc, "Ocena opisowa", v(data, "ocena_zakladowa_opis"))
    _sig_row(doc, "Data, pieczęć i podpis zakładowego opiekuna", v(data, "podpis_zakladowy"))

    _para(doc, "Ocena uczelnianego opiekuna praktyki", bold=True, size=FONT_BODY, space_before=4, space_after=2)
    _field_row_2col(doc, "Ocena parametryczna (skala 2–5)", v(data, "ocena_uczelniana_param") or "—",
                    "", "", w1=80, w2=80)
    _field_single(doc, "Ocena opisowa", v(data, "ocena_uczelniana_opis"))
    _sig_row(doc, "Data, pieczęć i podpis uczelnianego opiekuna", v(data, "podpis_uczelniany"))

    _para(doc, "Ocena sprawozdania z praktyki", bold=True, size=FONT_BODY, space_before=4, space_after=2)
    _field_row_2col(doc, "Ocena parametryczna (skala 2–5)", v(data, "ocena_sprawozdania") or "—",
                    "", "", w1=80, w2=80)
    _sig_row(doc, "Data i podpis uczelnianego opiekuna praktyki", v(data, "podpis_sprawozdanie"))

    return _to_bytes(doc)


# ─── Zał. 4 – Potwierdzenie efektów uczenia się ───────────────────────────────

def gen_zal4(data, effects=None):
    doc = _new_doc()
    _letterhead(doc, "4")
    _doc_title(doc, "Potwierdzenie uzyskania efektów uczenia się\nw ramach praktyki zawodowej")

    _field_row_2col(doc, "Student/ka", v(data, "imie_nazwisko"),
                    "Nr albumu", v(data, "nr_albumu"), w1=105, w2=55)
    _field_row_2col(doc, "Kierunek studiów", "Informatyka",
                    "Specjalność", v(data, "specjalnosc"), w1=55, w2=105)
    _field_single(doc, "Wymiar godzin", v(data, "wymiar_godzin") or "960")

    wym = v(data, "wymiar_godzin") or "960"
    _para(doc, f"W ramach praktyki zawodowej zrealizowanej w wymiarze {wym} godzin "
               "uzyskał/a lub nie uzyskał/a zakładane dla praktyki zawodowej efekty uczenia się:",
          size=FONT_BODY, space_before=4, space_after=4)

    # Tabela efektów
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    _cell_para(hdr[0], "Nr", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr[1], "Efekty uczenia się", bold=True, size=Pt(9))
    _cell_para(hdr[2], "Potwierdzenie uzyskania efektów", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in hdr:
        _set_cell_bg(c)

    efekty_map = {}
    for e in (data.get("efekty") or []):
        efekty_map[e["nr"]] = e.get("status", "")

    effect_list = effects or []
    for eff in effect_list:
        row = tbl.add_row()
        _cell_para(row.cells[0], f"{eff.nr:02d}", size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], eff.opis, size=Pt(9))
        status = efekty_map.get(eff.nr, "")
        _cell_para(row.cells[2], status or "uzyskał/a * nie uzyskał/a*",
                   size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_col_widths(tbl, [12, 110, 38])

    _section_label(doc, "Potwierdzenia")
    _sig_row(doc,
             "Potwierdzenie bezpośredniego opiekuna zakładowego", v(data, "potwierdzenie_opiekuna"),
             "Opinia opiekuna uczelnianego", v(data, "opinia_opiekuna"))

    _para(doc, "*) niepotrzebne skreślić", size=Pt(9), italic=True, space_before=4)

    return _to_bytes(doc)


# ─── Zał. 4a – Merytoryczna ocena wniosku studenta ───────────────────────────

def gen_zal4a(data, effects=None):
    doc = _new_doc()
    _letterhead(doc, "4a")
    _doc_title(doc, "Merytoryczna ocena wniosku studenta\no zaliczenie efektów uczenia się")

    _field_row_2col(doc, "Student/ka", v(data, "imie_nazwisko"),
                    "Nr albumu", v(data, "nr_albumu"), w1=105, w2=55)
    _field_single(doc, "Data złożenia wniosku", v(data, "data_zlozenia"))

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    _cell_para(hdr[0], "Nr", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr[1], "Efekty uczenia się", bold=True, size=Pt(9))
    _cell_para(hdr[2], "Zasadność", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr[3], "Uzasadnienie", bold=True, size=Pt(9))
    for c in hdr:
        _set_cell_bg(c)

    oceny_map = {}
    for e in (data.get("ocena_efektow") or []):
        oceny_map[e["nr"]] = e

    effect_list = effects or []
    for eff in effect_list:
        row = tbl.add_row()
        rec = oceny_map.get(eff.nr, {})
        _cell_para(row.cells[0], f"{eff.nr:02d}", size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], eff.opis, size=Pt(9))
        _cell_para(row.cells[2], rec.get("zasadny", "") or "—",
                   size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[3], rec.get("uzasadnienie", "") or "—", size=Pt(9))
    _set_col_widths(tbl, [12, 85, 25, 58])

    _field_single(doc, "Rekomendacja", v(data, "rekomendacja"))
    _field_single(doc, "Uwagi", v(data, "uwagi"))
    _sig_row(doc,
             "Data oceny", v(data, "data_oceny"),
             "Podpis uczelnianego opiekuna", v(data, "podpis_uopz"))

    return _to_bytes(doc)


# ─── Zał. 4b – Wniosek o zaliczenie efektów uczenia się ──────────────────────

def gen_zal4b(data, effects=None):
    doc = _new_doc()
    _letterhead(doc, "4b")
    _doc_title(doc, "Wniosek o zaliczenie efektów uczenia się dla praktyki zawodowej\nna podstawie pracy zawodowej, stażu lub prowadzenia działalności gospodarczej")

    _field_row_2col(doc, "Student/ka", v(data, "imie_nazwisko"),
                    "Nr albumu", v(data, "nr_albumu"), w1=105, w2=55)
    _field_row_2col(doc, "Kierunek studiów", "Informatyka",
                    "Specjalność", v(data, "specjalnosc"), w1=55, w2=105)
    _field_single(doc, "Pracodawca / podmiot", v(data, "pracodawca"))
    _field_single(doc, "Adres pracodawcy", v(data, "adres_pracodawcy"))
    _field_row_2col(doc, "Stanowisko / rodzaj działalności", v(data, "stanowisko"),
                    "", "", w1=160, w2=0)
    _field_row_2col(doc, "Okres zatrudnienia: od", v(data, "okres_od"),
                    "do", v(data, "okres_do"))

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    _cell_para(hdr[0], "Nr", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr[1], "Efekty uczenia się", bold=True, size=Pt(9))
    _cell_para(hdr[2], "Uzasadnienie", bold=True, size=Pt(9))
    _cell_para(hdr[3], "Dowody (dokumenty)", bold=True, size=Pt(9))
    for c in hdr:
        _set_cell_bg(c)

    efekty_map = {}
    for e in (data.get("efekty_wniosek") or []):
        efekty_map[e["nr"]] = e

    effect_list = effects or []
    for eff in effect_list:
        row = tbl.add_row()
        rec = efekty_map.get(eff.nr, {})
        _cell_para(row.cells[0], f"{eff.nr:02d}", size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], eff.opis, size=Pt(9))
        _cell_para(row.cells[2], rec.get("uzasadnienie", "") or "—", size=Pt(9))
        _cell_para(row.cells[3], rec.get("dowody", "") or "—", size=Pt(9))
    _set_col_widths(tbl, [12, 60, 60, 48])

    _field_single(doc, "Wykaz dokumentów potwierdzających", v(data, "wykaz_dokumentow"))
    _sig_row(doc,
             "Data", v(data, "data"),
             "Podpis studenta", v(data, "podpis_studenta"))

    return _to_bytes(doc)


# ─── Zał. 5 – Kwestionariusz ankiety ──────────────────────────────────────────

def gen_zal5(data):
    doc = _new_doc(top=25, bot=25, left=25, right=25)
    _letterhead(doc, "5")
    _doc_title(doc, "Kwestionariusz ankiety oceniający przebieg praktyk zawodowych\nrealizowanych w ramach programów studiów\nw Instytucie Informatyki Stosowanej ANS w Elblągu")

    _para(doc, "W trosce o stałe podnoszenie jakości przebiegu praktyk zawodowych zwracamy się do Pani/Pana z prośbą o wypełnienie anonimowej ankiety.",
          size=Pt(10), italic=True, space_before=2, space_after=4)

    # Metryczka
    tbl_m = doc.add_table(rows=5, cols=2)
    tbl_m.style = "Table Grid"
    _set_col_widths(tbl_m, [60, 100])
    rows_m = tbl_m.rows
    _cell_para(rows_m[0].cells[0], "Metryczka", bold=True, size=Pt(10))
    _cell_para(rows_m[0].cells[1], "", size=Pt(10))
    _cell_para(rows_m[1].cells[0], "Rok akademicki", size=Pt(10))
    _cell_para(rows_m[1].cells[1], v(data, "rok_akademicki"), size=Pt(10))
    _cell_para(rows_m[2].cells[0], "Kierunek studiów", size=Pt(10))
    _cell_para(rows_m[2].cells[1], "Informatyka", size=Pt(10))
    _cell_para(rows_m[3].cells[0], "Forma studiów", size=Pt(10))
    _cell_para(rows_m[3].cells[1], v(data, "forma_studiow") or "stacjonarne", size=Pt(10))
    _cell_para(rows_m[4].cells[0], "Semestr studiów", size=Pt(10))
    _cell_para(rows_m[4].cells[1], v(data, "semestr"), size=Pt(10))

    _para(doc, "Prosimy zaznaczyć przy każdym pytaniu X w wybranym polu odpowiedzi.",
          bold=True, size=Pt(10), space_before=6, space_after=4)

    OPTIONS = ["zdecydowanie tak", "raczej tak", "trudno powiedzieć", "raczej nie", "zdecydowanie nie"]

    pytania_map = {}
    for p in (data.get("pytania") or []):
        pytania_map[p["nr"]] = p.get("odpowiedz", "")

    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    _cell_para(hdr[0], "Nr", bold=True, size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr[1], "Pytanie", bold=True, size=Pt(8))
    for i, opt in enumerate(OPTIONS):
        _cell_para(hdr[2+i], opt, bold=True, size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in hdr:
        _set_cell_bg(c)
    _set_col_widths(tbl, [8, 75, 22, 18, 22, 18, 22])

    for qi, question in enumerate(SURVEY_QUESTIONS, 1):
        row = tbl.add_row()
        _cell_para(row.cells[0], str(qi), size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], question, size=Pt(8))
        ans = pytania_map.get(qi, "")
        for i, opt in enumerate(OPTIONS):
            mark = "X" if ans == opt else ""
            _cell_para(row.cells[2+i], mark, size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)

    _field_single(doc, "Dodatkowe uwagi dotyczące przebiegu praktyki zawodowej",
                  v(data, "uwagi"))

    _para(doc, "Dziękujemy za udział w badaniu", bold=True, size=Pt(10),
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8)

    return _to_bytes(doc)


# ─── Zał. 6 – Dziennik praktyki zawodowej ─────────────────────────────────────

def gen_zal6(data):
    doc = _new_doc(top=20, bot=15, left=25, right=25)
    _letterhead(doc, "6")
    _doc_title(doc, "Dziennik praktyki zawodowej")

    _field_row_2col(doc, "Student", v(data, "imie_nazwisko"),
                    "Nr albumu", v(data, "nr_albumu"), w1=110, w2=50)
    _field_row_2col(doc, "Kierunek", "Informatyka",
                    "Specjalność", v(data, "specjalnosc"), w1=55, w2=105)
    _field_row_2col(doc, "Forma studiów", v(data, "rodzaj_studiow") or "stacjonarne",
                    "Rok akademicki", v(data, "rok_akademicki"))
    _field_single(doc, "Miejsce odbywania praktyki", v(data, "miejsce_praktyki"))
    _field_row_2col(doc, "Data rozpoczęcia praktyki", v(data, "data_start"),
                    "Data zakończenia praktyki", v(data, "data_end"))

    # Dziennik entries
    dziennik = data.get("dziennik") or []
    if dziennik:
        _section_label(doc, "Wpisy dziennika")
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        _cell_para(hdr[0], "Dzień", bold=True, size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(hdr[1], "Data", bold=True, size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(hdr[2], "Opis czynności", bold=True, size=Pt(8))
        _cell_para(hdr[3], "Efekty", bold=True, size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(hdr[4], "Podpis opiekuna", bold=True, size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
        for c in hdr:
            _set_cell_bg(c)
        _set_col_widths(tbl, [15, 22, 85, 20, 28])

        for entry in dziennik:
            row = tbl.add_row()
            _cell_para(row.cells[0], str(entry.get("dzien", "")), size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_para(row.cells[1], entry.get("data", ""), size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_para(row.cells[2], entry.get("opis", ""), size=Pt(9))
            _cell_para(row.cells[3], entry.get("efekty", ""), size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_para(row.cells[4], entry.get("podpis", ""), size=Pt(8), align=WD_ALIGN_PARAGRAPH.CENTER)

    _field_single(doc, "Wykaz załączników", v(data, "wykaz_zalacznikow"))

    return _to_bytes(doc)


# ─── Zał. 7 – Sprawozdanie (stacjonarne) ──────────────────────────────────────

def gen_zal7(data):
    doc = _new_doc()
    _letterhead(doc, "7")

    _field_row_2col(doc, "Student", v(data, "imie_nazwisko"),
                    "Nr albumu", v(data, "nr_albumu"), w1=110, w2=50)
    _field_row_2col(doc, "Kierunek", "Informatyka",
                    "Specjalność", v(data, "specjalnosc"), w1=55, w2=105)
    _field_row_2col(doc, "Studia", v(data, "rodzaj_studiow") or "inżynierskie stacjonarne",
                    "Rok akademicki", v(data, "rok_akademicki"))

    _para(doc, "", space_before=0, space_after=6)
    _doc_title(doc, "Sprawozdanie studenta\nz praktyki zawodowej")
    p = _para(doc, "odbytej w  ", space_before=2, space_after=6)
    _run(p, v(data, "miejsce_praktyki") or "……………………………………………………", bold=True)

    _section_label(doc, "1. Charakterystyka miejsca odbywania praktyki")
    _para(doc, "(Krótki opis instytucji, w której odbywała się praktyka zawodowa)", italic=True, size=Pt(10), space_after=4)
    txt1 = v(data, "charakterystyka")
    if txt1:
        _para(doc, txt1, size=FONT_BODY, space_after=6)
    else:
        for _ in range(8):
            _para(doc, "", space_after=10)

    _section_label(doc, "2. Opis i analiza wykonywanych prac")
    _para(doc, "(Syntetyczny opis wykonanych prac)", italic=True, size=Pt(10), space_after=4)
    txt2 = v(data, "opis_prac")
    if txt2:
        _para(doc, txt2, size=FONT_BODY, space_after=6)
    else:
        for _ in range(10):
            _para(doc, "", space_after=10)

    _section_label(doc, "3. Wiedza i umiejętności uzyskane w trakcie praktyki")
    _para(doc, "(Samoocena w zakresie nabytych kompetencji oraz osiągniętych efektów uczenia się)", italic=True, size=Pt(10), space_after=4)
    txt3 = v(data, "wiedza_umiejetnosci")
    if txt3:
        _para(doc, txt3, size=FONT_BODY, space_after=6)
    else:
        for _ in range(8):
            _para(doc, "", space_after=10)

    _sig_row(doc,
             "Data sporządzenia sprawozdania", v(data, "data"),
             "Podpis studenta", v(data, "podpis_studenta"))

    return _to_bytes(doc)


# ─── Zał. 7a – Sprawozdanie (niestacjonarne) ──────────────────────────────────

def gen_zal7a(data):
    doc = _new_doc()
    _letterhead(doc, "7a")

    _field_row_2col(doc, "Student", v(data, "imie_nazwisko"),
                    "Nr albumu", v(data, "nr_albumu"), w1=110, w2=50)
    _field_row_2col(doc, "Kierunek", "Informatyka",
                    "Specjalność", v(data, "specjalnosc"), w1=55, w2=105)
    _field_row_2col(doc, "Studia", "inżynierskie niestacjonarne",
                    "Rok akademicki", v(data, "rok_akademicki"))

    _doc_title(doc, "Sprawozdanie studenta\nz praktyki zawodowej realizowanej\nna podstawie pracy zawodowej lub działalności gospodarczej")
    p = _para(doc, "odbytej w  ", space_before=2, space_after=6)
    _run(p, v(data, "miejsce_praktyki") or "……………………………………………………", bold=True)

    _section_label(doc, "1. Charakterystyka miejsca pracy")
    _para(doc, "(Krótki opis instytucji)", italic=True, size=Pt(10), space_after=4)
    txt1 = v(data, "charakterystyka")
    _para(doc, txt1 or "", size=FONT_BODY, space_after=6) if txt1 else [_para(doc, "", space_after=10) for _ in range(8)]

    _section_label(doc, "2. Opis i analiza wykonywanych prac")
    _para(doc, "(Syntetyczny opis wykonanych prac)", italic=True, size=Pt(10), space_after=4)
    txt2 = v(data, "opis_prac")
    _para(doc, txt2 or "", size=FONT_BODY, space_after=6) if txt2 else [_para(doc, "", space_after=10) for _ in range(10)]

    _section_label(doc, "3. Wiedza i umiejętności uzyskane w trakcie pracy zawodowej lub działalności gospodarczej")
    _para(doc, "(Samoocena w zakresie nabytych kompetencji oraz osiągniętych efektów uczenia się)", italic=True, size=Pt(10), space_after=4)
    txt3 = v(data, "wiedza_umiejetnosci")
    _para(doc, txt3 or "", size=FONT_BODY, space_after=6) if txt3 else [_para(doc, "", space_after=10) for _ in range(8)]

    _sig_row(doc,
             "Data i podpis studenta", v(data, "podpis_studenta"),
             "Data i podpis bezpośredniego przełożonego lub uczelnianego opiekuna praktyk",
             v(data, "podpis_przelozonego"))

    return _to_bytes(doc)


# ─── Zał. 8 – Protokół zaliczenia praktyki zawodowej ─────────────────────────

def gen_zal8(data):
    doc = _new_doc(top=15, bot=15, left=20, right=20)
    _letterhead(doc, "8")

    _field_row_2col(doc, "Student", v(data, "imie_nazwisko"),
                    "Nr albumu", v(data, "nr_albumu"), w1=115, w2=55)

    _doc_title(doc, "Protokół zaliczenia praktyki zawodowej (PZ)")

    # Miejsca praktyki
    _section_label(doc, "Miejsce i okres realizacji praktyki zawodowej")
    miejsca = data.get("miejsca_praktyki") or []
    if miejsca:
        tbl_m = doc.add_table(rows=1, cols=3)
        tbl_m.style = "Table Grid"
        hdr = tbl_m.rows[0].cells
        _cell_para(hdr[0], "Miejsce odbywania praktyki", bold=True, size=Pt(9))
        _cell_para(hdr[1], "Okres", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(hdr[2], "Liczba dni", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        for c in hdr:
            _set_cell_bg(c)
        for m in miejsca:
            row = tbl_m.add_row()
            _cell_para(row.cells[0], m.get("nazwa", ""), size=Pt(9))
            _cell_para(row.cells[1], m.get("okres", ""), size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_para(row.cells[2], str(m.get("dni", "")), size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_col_widths(tbl_m, [110, 40, 30])
    _para(doc, "", space_before=0, space_after=3)

    # Oceny
    _section_label(doc, "Oceny składowe")
    _field_row_2col(doc,
                    "Ocena za sprawozdanie z praktyki (S)", v(data, "ocena_s") or "—",
                    "Data sprawozdania", v(data, "data_s") or "—")
    _field_row_2col(doc,
                    "Ocena uczelnianego opiekuna (U)", v(data, "ocena_u") or "—",
                    "Ocena zakładowego opiekuna (Z)", v(data, "ocena_z") or "—")

    # Mini-zadania
    _section_label(doc, "Mini-zadania zawodowe komisji egzaminacyjnej")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    _cell_para(hdr[0], "Lp.", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr[1], "Pytania / mini zadania zawodowe", bold=True, size=Pt(9))
    _cell_para(hdr[2], "Oceny cząstkowe (skala 2–5)", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr[3], "", size=Pt(9))
    for c in hdr:
        _set_cell_bg(c)

    mini = data.get("mini_zadania") or [{"tresc": "", "ocena": ""} for _ in range(3)]
    for i, m in enumerate(mini, 1):
        row = tbl.add_row()
        _cell_para(row.cells[0], str(i), size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], m.get("tresc", ""), size=Pt(9))
        _cell_para(row.cells[2], m.get("ocena", ""), size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[3], "", size=Pt(9))

    row_e = tbl.add_row()
    _cell_para(row_e.cells[0], "", size=Pt(9))
    _cell_para(row_e.cells[1], "Łączna ocena za mini zadania (E) = średnia arytmetyczna ocen cząstkowych", bold=True, size=Pt(9))
    _cell_para(row_e.cells[2], v(data, "ocena_e") or "—", bold=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(row_e.cells[3], "", size=Pt(9))
    _set_col_widths(tbl, [12, 110, 30, 28])

    _para(doc, "", space_before=4, space_after=2)
    ocena_k = v(data, "ocena_k") or "—"
    _para(doc,
          f"Ocena końcowa za PZ:  0,4·E + 0,1·S + 0,2·U + 0,3·Z  =  K  =  {ocena_k}",
          bold=True, size=FONT_BODY, space_before=4, space_after=4)

    # Komisja
    _section_label(doc, "Skład komisji")
    _field_row_2col(doc, "Data zaliczenia", v(data, "data_zaliczenia"),
                    "Skład komisji", v(data, "sklad_komisji"))
    _field_row_2col(doc, "Przewodniczący Komisji", v(data, "przewodniczacy"),
                    "Członek komisji", v(data, "czlonek_2"))
    if v(data, "czlonek_3"):
        _field_single(doc, "Członek komisji", v(data, "czlonek_3"))

    _para(doc, f"Zaliczam praktykę zawodową na ocenę (K): {ocena_k}",
          bold=True, size=FONT_BODY, space_before=6, space_after=4)

    _para(doc, "*) U i Z – oceny odpowiednio Opiekuna uczelnianego i Opiekuna zakładowego zapisane w Karcie praktyki zawodowej.",
          italic=True, size=Pt(9), space_before=4)

    return _to_bytes(doc)


# ─── Zał. 9 – Oświadczenie instytucji ─────────────────────────────────────────

def gen_zal9(data):
    doc = _new_doc(top=25, bot=25, left=25, right=25)

    # Brak letterheadu – oryginał zaczyna się od miejscowości i pieczęci
    p_city = _para(doc, "", space_before=0, space_after=4)
    _run(p_city, v(data, "miejscowosc") or "……………………………………", bold=False)
    _run(p_city, ",  dnia  ")
    _run(p_city, v(data, "data") or "………………………………", bold=False)
    p_city.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _para(doc, "(miejscowość)", italic=True, size=Pt(9), space_before=0, space_after=10)
    _para(doc, "………………………………………………………", size=Pt(9), space_before=0, space_after=2)
    _para(doc, "Pieczęć firmy", italic=True, size=Pt(9), space_before=0, space_after=14)

    _doc_title(doc, "Oświadczenie instytucji\nw sprawie przyjęcia studenta na praktykę zawodową",
               center=True)

    p_intro = _para(doc, "W imieniu  ", space_before=4, space_after=2)
    _run(p_intro, v(data, "nazwa_instytucji") or "……………………………………………………………………………………", bold=True)
    _para(doc, "(nazwa instytucji)", italic=True, size=Pt(9), space_after=4)

    p_decl = _para(doc, "oświadczam, że w terminie od  ", space_before=2, space_after=2)
    _run(p_decl, v(data, "termin_od") or "……………………", bold=True)
    _run(p_decl, "  do  ")
    _run(p_decl, v(data, "termin_do") or "……………………", bold=True)
    _run(p_decl, "  przyjmiemy na praktykę zawodową studenta Instytutu Informatyki Stosowanej im. Krzysztofa Brzeskiego")

    p_student = _para(doc, v(data, "imie_nazwisko") or "……………………………………………………………", bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)
    _para(doc, "(imię i nazwisko studenta)", italic=True, size=Pt(9),
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    p_ucz = _para(doc, "Akademii Nauk Stosowanych w Elblągu,  kierunek: informatyka,  rok studiów: IV,  nr albumu: ", space_before=2, space_after=4)
    _run(p_ucz, v(data, "nr_albumu") or "……………", bold=True)

    _para(doc, "Zakładowym opiekunem praktyki będzie:", size=FONT_BODY, space_before=6, space_after=2)
    _field_single(doc, "(imię i nazwisko, stanowisko)", v(data, "opiekun_imie_nazwisko"))
    _field_row_2col(doc, "Telefon", v(data, "opiekun_telefon"),
                    "E-mail", v(data, "opiekun_email"))

    _para(doc, "Osobą upoważnioną do podpisania porozumienia dotyczącego prowadzenia praktyki zawodowej jest ze strony naszej instytucji",
          size=FONT_BODY, space_before=6, space_after=2)
    _field_single(doc, "(imię i nazwisko, stanowisko)", v(data, "upowazniont_imie_nazwisko"))

    _para(doc, "", space_before=12, space_after=2)
    _sig_row(doc, "Pieczęć imienna i podpis", v(data, "podpis"))

    return _to_bytes(doc)


# ─── dispatcher ───────────────────────────────────────────────────────────────

def generate(zal_key, data, effects=None):
    """Generate docx for given attachment key. Returns BytesIO buffer."""
    generators = {
        "zal1":  lambda: gen_zal1(data),
        "zal2":  lambda: gen_zal2(data),
        "zal2a": lambda: gen_zal2a(data, effects),
        "zal3":  lambda: gen_zal3(data),
        "zal4":  lambda: gen_zal4(data, effects),
        "zal4a": lambda: gen_zal4a(data, effects),
        "zal4b": lambda: gen_zal4b(data, effects),
        "zal5":  lambda: gen_zal5(data),
        "zal6":  lambda: gen_zal6(data),
        "zal7":  lambda: gen_zal7(data),
        "zal7a": lambda: gen_zal7a(data),
        "zal8":  lambda: gen_zal8(data),
        "zal9":  lambda: gen_zal9(data),
    }
    fn = generators.get(zal_key)
    if fn is None:
        raise ValueError(f"Unknown attachment key: {zal_key}")
    return fn()
